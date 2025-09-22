"""
Document Processor Module

Handles document file discovery, content extraction, and preprocessing.
Supports multiple document formats including text, PDF, Word documents.
"""

import os
import logging
from pathlib import Path
from typing import List, Optional, Dict
import mimetypes


class DocumentProcessor:
    """Handles document file processing and content extraction."""
    
    def __init__(self, config: Dict):
        """Initialize document processor with configuration."""
        self.config = config
        self.logger = logging.getLogger(__name__)
        self.supported_formats = config.get('supported_formats', ['.txt', '.md'])
        self.max_file_size_mb = config.get('max_file_size_mb', 50)
    
    def get_document_files(self, input_folder: Path) -> List[Path]:
        """Get list of supported document files from input folder."""
        document_files = []
        
        if not input_folder.exists():
            self.logger.error(f"Input folder does not exist: {input_folder}")
            return document_files
        
        for file_path in input_folder.iterdir():
            if file_path.is_file() and self._is_supported_file(file_path):
                if self._check_file_size(file_path):
                    document_files.append(file_path)
                else:
                    self.logger.warning(f"File too large, skipping: {file_path.name}")
        
        self.logger.info(f"Found {len(document_files)} supported document files")
        return sorted(document_files)
    
    def extract_content(self, file_path: Path) -> str:
        """Extract text content from document file."""
        try:
            content = ""
            file_extension = file_path.suffix.lower()
            
            if file_extension in ['.txt', '.md']:
                content = self._extract_text_file(file_path)
            elif file_extension == '.pdf':
                content = self._extract_pdf_content(file_path)
            elif file_extension in ['.docx', '.doc']:
                content = self._extract_word_content(file_path)
            else:
                self.logger.warning(f"Unsupported file format: {file_extension}")
                return ""
            
            # Clean and preprocess content
            content = self._preprocess_content(content)
            
            self.logger.debug(f"Extracted {len(content)} characters from {file_path.name}")
            return content
            
        except Exception as e:
            self.logger.error(f"Error extracting content from {file_path.name}: {e}")
            return ""
    
    def _is_supported_file(self, file_path: Path) -> bool:
        """Check if file format is supported."""
        return file_path.suffix.lower() in self.supported_formats
    
    def _check_file_size(self, file_path: Path) -> bool:
        """Check if file size is within limits."""
        try:
            file_size_mb = file_path.stat().st_size / (1024 * 1024)
            return file_size_mb <= self.max_file_size_mb
        except Exception as e:
            self.logger.error(f"Error checking file size for {file_path.name}: {e}")
            return False
    
    def _extract_text_file(self, file_path: Path) -> str:
        """Extract content from plain text or markdown file."""
        encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
        
        for encoding in encodings:
            try:
                with open(file_path, 'r', encoding=encoding) as file:
                    content = file.read()
                return content
            except (UnicodeDecodeError, UnicodeError):
                continue
            except Exception as e:
                self.logger.error(f"Error reading text file {file_path.name}: {e}")
                break
        
        self.logger.error(f"Unable to read text file with any supported encoding: {file_path.name}")
        return ""
    
    def _extract_pdf_content(self, file_path: Path) -> str:
        """Extract content from PDF file."""
        try:
            # Try to import PDF libraries
            try:
                import PyPDF2
                return self._extract_pdf_pypdf2(file_path)
            except ImportError:
                try:
                    import pdfplumber
                    return self._extract_pdf_pdfplumber(file_path)
                except ImportError:
                    self.logger.error("PDF processing requires PyPDF2 or pdfplumber. Install with: pip install PyPDF2 pdfplumber")
                    return ""
        except Exception as e:
            self.logger.error(f"Error extracting PDF content from {file_path.name}: {e}")
            return ""
    
    def _extract_pdf_pypdf2(self, file_path: Path) -> str:
        """Extract PDF content using PyPDF2."""
        import PyPDF2
        
        content = ""
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            for page in reader.pages:
                content += page.extract_text() + "\n"
        
        return content
    
    def _extract_pdf_pdfplumber(self, file_path: Path) -> str:
        """Extract PDF content using pdfplumber."""
        import pdfplumber
        
        content = ""
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    content += page_text + "\n"
        
        return content
    
    def _extract_word_content(self, file_path: Path) -> str:
        """Extract content from Word document."""
        try:
            # Try to import Word document libraries
            try:
                import docx
                return self._extract_docx_content(file_path)
            except ImportError:
                try:
                    import python_docx
                    return self._extract_docx_content_alt(file_path)
                except ImportError:
                    self.logger.error("Word document processing requires python-docx. Install with: pip install python-docx")
                    return ""
        except Exception as e:
            self.logger.error(f"Error extracting Word content from {file_path.name}: {e}")
            return ""
    
    def _extract_docx_content(self, file_path: Path) -> str:
        """Extract DOCX content using python-docx."""
        import docx
        
        doc = docx.Document(file_path)
        content = ""
        
        for paragraph in doc.paragraphs:
            content += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    content += cell.text + " "
                content += "\n"
        
        return content
    
    def _extract_docx_content_alt(self, file_path: Path) -> str:
        """Alternative method for extracting DOCX content."""
        # Fallback implementation if needed
        return ""
    
    def _preprocess_content(self, content: str) -> str:
        """Clean and preprocess extracted content."""
        if not content:
            return ""
        
        # Remove excessive whitespace
        lines = [line.strip() for line in content.split('\n')]
        lines = [line for line in lines if line]  # Remove empty lines
        
        # Join with single spaces
        content = '\n'.join(lines)
        
        # Limit content length to avoid token limits
        max_content_length = 50000  # Reasonable limit for LLM processing
        if len(content) > max_content_length:
            content = content[:max_content_length] + "\n[Content truncated due to length]"
            self.logger.info("Content truncated due to length limit")
        
        return content
    
    def get_file_metadata(self, file_path: Path) -> Dict:
        """Extract metadata from file."""
        try:
            stat = file_path.stat()
            
            metadata = {
                'filename': file_path.name,
                'file_size': stat.st_size,
                'created_time': stat.st_ctime,
                'modified_time': stat.st_mtime,
                'file_extension': file_path.suffix.lower(),
                'mime_type': mimetypes.guess_type(file_path)[0]
            }
            
            return metadata
            
        except Exception as e:
            self.logger.error(f"Error getting file metadata for {file_path.name}: {e}")
            return {'filename': file_path.name}
    
    def validate_input_folder(self, folder_path: Path) -> bool:
        """Validate that input folder exists and is accessible."""
        if not folder_path.exists():
            self.logger.error(f"Input folder does not exist: {folder_path}")
            return False
        
        if not folder_path.is_dir():
            self.logger.error(f"Input path is not a directory: {folder_path}")
            return False
        
        try:
            # Test read access
            list(folder_path.iterdir())
            return True
        except PermissionError:
            self.logger.error(f"Permission denied accessing folder: {folder_path}")
            return False
        except Exception as e:
            self.logger.error(f"Error accessing folder {folder_path}: {e}")
            return False
    
    # Add this method to the DocumentProcessor class
    
    def chunk_document(self, content: str) -> List[str]:
        """Split document content into overlapping chunks for RAG processing."""
        try:
            chunk_size = self.config.get('chunk_size', 1000)
            chunk_overlap = self.config.get('chunk_overlap', 200)
            
            # Simple chunking by characters with overlap
            chunks = []
            start = 0
            content_length = len(content)
            
            while start < content_length:
                end = min(start + chunk_size, content_length)
                chunks.append(content[start:end])
                start = end - chunk_overlap
            
            self.logger.debug(f"Created {len(chunks)} chunks from document")
            return chunks
        except Exception as e:
            self.logger.error(f"Error chunking document: {e}")
            return [content]  # Return full content as single chunk on error